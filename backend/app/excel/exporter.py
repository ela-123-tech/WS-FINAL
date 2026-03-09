from __future__ import annotations
import os
from datetime import datetime
from docx import Document
from app import models

EXPORT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "exports"))
os.makedirs(EXPORT_DIR, exist_ok=True)


def _format_value(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return ", ".join(str(x) for x in value)
    return str(value)


def _write_table(document: Document, headers: list[str], rows: list[list[object]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx >= len(cells):
                break
            cells[idx].text = _format_value(value)
    document.add_paragraph()
    return table


def export_all(trip: models.Trip) -> str:
    doc = Document()
    doc.add_heading("Trip Summary", level=1)
    summary_rows = [
        ["Origin", trip.origin_display],
        ["Destination", trip.destination_display],
        ["Created", trip.created_at.isoformat()],
    ]
    if trip.route:
        summary_rows.extend([
            ["Distance (km)", round(trip.route.distance_m / 1000, 2)],
            ["Duration (min)", round(trip.route.duration_s / 60, 1)],
        ])
    _write_table(doc, ["Field", "Value"], summary_rows)

    doc.add_heading("Transport Options", level=2)
    transport_headers = ["ID", "Mode", "Operator", "Title", "Departure", "Arrival", "Duration (min)",
                         "Distance (km)", "Price (INR)", "Stops", "Booking URL"]
    transport_rows = []
    for t in trip.transport_options:
        transport_rows.append([
            t.id,
            t.mode,
            t.operator,
            t.title,
            t.departure_iso,
            t.arrival_iso,
            round((t.duration_s or 0) / 60, 1) if t.duration_s else "",
            round((t.distance_m or 0) / 1000, 2) if t.distance_m else "",
            t.price_inr,
            t.stops,
            t.booking_url,
        ])
    if transport_rows:
        _write_table(doc, transport_headers, transport_rows)

    def add_poi_section(title: str, kind: str):
        items = [p for p in trip.pois if p.kind == kind]
        if not items:
            return
        doc.add_heading(title, level=3)
        poi_headers = ["ID", "Name", "Lat", "Lng", "Address", "Tags", "Source"]
        poi_rows = []
        for p in items:
            poi_rows.append([
                p.id,
                p.name,
                p.lat,
                p.lng,
                p.address,
                ", ".join(p.tags or []),
                p.source,
            ])
        _write_table(doc, poi_headers, poi_rows)

    add_poi_section("Hotels", "hotel")
    add_poi_section("Food", "food")
    add_poi_section("Attractions", "attraction")

    if trip.route and trip.route.steps:
        doc.add_heading("Route Steps", level=2)
        route_headers = ["#", "Name", "Instruction", "Distance (m)", "Duration (s)"]
        route_rows = []
        for idx, step in enumerate(trip.route.steps or [], start=1):
            route_rows.append([
                idx,
                step.get("name"),
                step.get("instruction"),
                step.get("distance"),
                step.get("duration"),
            ])
        _write_table(doc, route_headers, route_rows)

    meta = trip.meta or {}
    weather = meta.get("weather") or {}
    if weather:
        doc.add_heading("Weather", level=2)
        weather_rows = []
        if weather.get("timezone"):
            weather_rows.append(["Timezone", weather.get("timezone")])
        if weather_rows:
            _write_table(doc, ["Field", "Value"], weather_rows)
        daily = weather.get("daily") or {}
        times = daily.get("time") or []
        tmax = daily.get("temperature_2m_max") or []
        tmin = daily.get("temperature_2m_min") or []
        pr = daily.get("precipitation_sum") or []
        daily_rows = []
        for i in range(min(len(times), len(tmax), len(tmin), len(pr))):
            daily_rows.append([times[i], tmax[i], tmin[i], pr[i]])
        if daily_rows:
            doc.add_heading("Daily Forecast", level=3)
            _write_table(doc, ["Date", "Tmax", "Tmin", "Precipitation"], daily_rows)

    dest = meta.get("destination_summary") or {}
    dest_items = []
    for key in ("title", "extract", "thumbnail", "page"):
        if dest.get(key):
            dest_items.append([key, dest.get(key)])
    if dest_items:
        doc.add_heading("Destination Information", level=2)
        _write_table(doc, ["Field", "Value"], dest_items)

    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"trip_{trip.id}_all_{ts}.docx"
    doc.save(os.path.join(EXPORT_DIR, filename))
    return filename


def export_final(trip: models.Trip, selection: models.TripSelection) -> str:
    doc = Document()
    doc.add_heading("Final Itinerary", level=1)
    summary_rows = [
        ["Origin", trip.origin_display],
        ["Destination", trip.destination_display],
        ["Created", selection.created_at.isoformat()],
    ]
    _write_table(doc, ["Field", "Value"], summary_rows)

    transport = next((t for t in trip.transport_options if t.id == selection.selected_transport_option_id), None)
    plan = next((p for p in trip.plans if p.id == selection.selected_plan_id), None)

    doc.add_heading("Selected Transport", level=2)
    if transport:
        transport_rows = [
            ["Mode", transport.mode],
            ["Title", transport.title],
            ["Operator", transport.operator],
            ["Departure", transport.departure_iso],
            ["Arrival", transport.arrival_iso],
            ["Duration (min)", round((transport.duration_s or 0) / 60, 1) if transport.duration_s else ""],
            ["Price (INR)", transport.price_inr],
        ]
        _write_table(doc, ["Field", "Value"], transport_rows)
    else:
        doc.add_paragraph("No transport option selected.")

    doc.add_heading("Selected Plan", level=2)
    if plan:
        plan_rows = [
            ["Plan", plan.name],
            ["Summary", plan.summary],
            ["Hotel POI ID", plan.hotel_poi_id],
            ["Food POI ID", plan.food_poi_id],
        ]
        _write_table(doc, ["Field", "Value"], plan_rows)
    else:
        doc.add_paragraph("No trip plan selected.")

    if plan and plan.timeline:
        doc.add_heading("Timeline", level=2)
        timeline_rows = []
        for row in plan.timeline:
            timeline_rows.append([row.get("t"), row.get("note")])
        if timeline_rows:
            _write_table(doc, ["Step", "Note"], timeline_rows)

    links = []
    if transport and transport.booking_url:
        links.append(["Booking URL", transport.booking_url])
    dest_page = (trip.meta or {}).get("destination_summary", {}).get("page")
    if dest_page:
        links.append(["Destination page", dest_page])
    if links:
        doc.add_heading("Links", level=2)
        _write_table(doc, ["Type", "Value"], links)

    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"trip_{trip.id}_final_{ts}.docx"
    doc.save(os.path.join(EXPORT_DIR, filename))
    return filename
